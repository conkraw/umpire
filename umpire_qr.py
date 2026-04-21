import streamlit as st
import pandas as pd
import qrcode
import io
import zipfile
import re

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from pathlib import Path

LOGO_PATH = Path("hershey_baseball_umpire_assessment.png")

# -----------------------------------
# PAGE SETUP
# -----------------------------------
st.set_page_config(page_title="Umpire QR Document Generator", layout="centered")

st.title("Umpire QR Document Generator")

REDCAP_LINK = "https://redcap.ctsi.psu.edu/surveys/?s=J3K9DE4TL8FR8XXC"

# -----------------------------------
# HELPERS
# -----------------------------------
def clean_filename(text):
    """Remove characters that are invalid in filenames."""
    text = str(text).strip()
    text = re.sub(r'[\\/*?:"<>|]', "", text)
    text = text.replace(" ", "_")
    return text

def format_date(date_val):
    """Convert Excel date into a readable string."""
    if pd.notna(date_val):
        try:
            return pd.to_datetime(date_val).strftime("%Y-%m-%d")
        except Exception:
            return str(date_val)
    return ""

def make_qr_image(data_text):
    """Create QR code image in memory."""
    qr = qrcode.QRCode(
        version=1,
        box_size=10,
        border=4
    )
    qr.add_data(data_text)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)
    return img_bytes

def make_docx(row):
    """Create one DOCX file for a single row."""
    coach = str(row["coach"]).strip()
    umpire = str(row["umpire"]).strip()
    date_str = format_date(row["date"])

    from urllib.parse import quote

    coach_encoded = quote(coach)
    umpire_encoded = quote(umpire)
    date_encoded = quote(date_str)
    
    qr_text = f"https://redcap.ctsi.psu.edu/surveys/?s=J3K9DE4TL8FR8XXC&coach={coach_encoded}&umpire={umpire_encoded}&date={date_encoded}"
    
    qr_img = make_qr_image(qr_text)

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Umpire Evaluation")
    run.bold = True

    doc.add_paragraph("")

    # Coach
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r1 = p1.add_run("Coach: ")
    r1.bold = True
    p1.add_run(coach)

    # Umpire
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = p2.add_run("Umpire: ")
    r2.bold = True
    p2.add_run(umpire)

    # Date
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r3 = p3.add_run("Date: ")
    r3.bold = True
    p3.add_run(date_str)

    doc.add_paragraph("")

    # Instructions
    p4 = doc.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p4.add_run("Please scan the QR code below to complete the feedback form.")

    doc.add_paragraph("")

    # QR code
    p5 = doc.add_paragraph()
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    qr_run = p5.add_run()
    qr_run.add_picture(qr_img, width=Inches(2.25))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# -----------------------------------
# FILE UPLOAD
# -----------------------------------
uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx"])

if uploaded_file is not None:
    try:
        df = pd.read_excel(uploaded_file)

        st.subheader("Preview of uploaded file")
        st.dataframe(df)

        required_cols = {"date", "coach", "umpire"}
        missing_cols = required_cols - set(df.columns)

        if missing_cols:
            st.error(f"Missing required column(s): {', '.join(missing_cols)}")
            st.stop()

        df = df.dropna(subset=["coach", "umpire"]).copy()

        if df.empty:
            st.warning("No usable rows found after removing blank coach/umpire rows.")
            st.stop()

        st.success(f"Found {len(df)} row(s) ready for document creation.")

        if st.button("Generate DOCX Files"):
            zip_buffer = io.BytesIO()

            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, row in df.iterrows():
                    coach = clean_filename(row["coach"])
                    umpire = clean_filename(row["umpire"])

                    filename = f"{coach}_{umpire}.docx"
                    docx_file = make_docx(row)

                    zf.writestr(filename, docx_file.read())

            zip_buffer.seek(0)

            st.success("Documents generated successfully.")

            st.download_button(
                label="Download ZIP File",
                data=zip_buffer,
                file_name="umpire_qr_documents.zip",
                mime="application/zip"
            )

    except Exception as e:
        st.error(f"Error reading file: {e}")
